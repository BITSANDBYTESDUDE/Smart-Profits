from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0B, 0x11, 0x20)
CYAN = RGBColor(0x0F, 0x9E, 0x94)
GOLD = RGBColor(0xC4, 0xA0, 0x35)
MUTED = RGBColor(0x5B, 0x6B, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def set_run(run, size=11, bold=False, color=BLACK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D8E0EA")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading_bar(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "C4A035")
    pBdr.append(left)
    pPr.append(pBdr)
    set_run(p.add_run(text), 13, True, NAVY)


def add_body(text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), size, False, BLACK)


def add_sub(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), 12, True, CYAN)


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75)
    set_run(p.add_run(text), 11, False, BLACK)


def kv_table(rows):
    t = doc.add_table(rows=len(rows), cols=2)
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text = ""
        c1.text = ""
        set_run(c0.paragraphs[0].add_run(k), 10, True, WHITE)
        set_run(c1.paragraphs[0].add_run(v), 10, False, BLACK)
        shade_cell(c0, "0F9E94")
        shade_cell(c1, "F3F6FB" if i % 2 == 0 else "FFFFFF")
        set_cell_border(c0)
        set_cell_border(c1)
        c0.width = Cm(5.4)
        c1.width = Cm(11.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
set_run(p.add_run("GSMA MENA IGNITE OPEN GATEWAY HACKATHON"), 10, True, CYAN)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
set_run(
    p.add_run("Nokia Network-as-Code  |  CAMARA APIs  |  Phase 1: Idea Submission"),
    10,
    False,
    MUTED,
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
set_run(p.add_run("IDEA CAPTURE TEMPLATE"), 11, True, GOLD)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
set_run(
    p.add_run(
        "Smart Profits: AI-Driven Financial Analytics & Telco-Secured Merchant Platform"
    ),
    18,
    True,
    NAVY,
)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
set_run(p.add_run("Phase 1 Idea Capture Document"), 12, False, MUTED)

add_heading_bar("1. Idea name and submission date")
kv_table(
    [
        (
            "Idea name",
            "Smart Profits: AI-Driven Financial Analytics & Telco-Secured Merchant Platform",
        ),
        ("Short name", "Smart Profits"),
        ("Submission date", "14 August 2026"),
        ("Phase 1 deadline", "23 August 2026"),
        ("Stage", "Phase 1 — Idea Submission (Idea Capture Template + Pitch Deck)"),
    ]
)

add_heading_bar("2. Applicant details")
kv_table(
    [
        ("Full name", "Israa Nael Hamad"),
        ("Team name", "Smart Profits Team"),
        (
            "Role",
            "Founder / Product Lead  |  Software Engineering Student (final year)",
        ),
        (
            "University / major",
            "University of Palestine (Gaza) — B.Sc. Software Engineering, final year",
        ),
        ("Location", "Gaza, Palestine — MENA"),
        ("Email", "tsraathmd@gmail.com"),
        ("Languages", "Arabic and English (RTL / LTR product UI)"),
    ]
)

add_heading_bar("3. Idea Summary")

add_sub("The problem")
add_body(
    "Independent merchants across the Middle East and North Africa run their businesses from messy Excel, CSV, PDF and image files, in Arabic and English. They have no CFO, and the platform that holds their profits and sales files has no bank-grade protection."
)
add_body(
    "Three risks sit on the same login. (1) Merchant scatter and financial blindness: sales files hide phantom profit because rent, salaries and utilities are missing, while wrong prices and dead stock drain cash. (2) Profit leakage: weak margins and unexamined SKUs. (3) Fraud and account takeover: a SIM swap defeats SMS OTP, so an attacker can open the advisor, download P&L reports, change prices, or steal the store's sales files. A session from another country uploading the monthly file can be a hijacked account."
)

add_sub("The innovative solution")
add_body(
    "Smart Profits is a bilingual merchant platform that turns a messy workbook into a store diagnosis (health, profit leaks, inventory, 30-day actions) and answers questions from the currently open file (for example: highest profit product)."
)
add_body(
    "What makes it a hackathon FinTech — not a generic Excel AI — is the Smart Guard AI Agent. On login, password reset, file upload, report export and sensitive price changes, the agent calls GSMA CAMARA APIs through Nokia Network-as-Code (Number Verification, SIM Swap, Location Verification). It then decides: Allow, Step-up, or Freeze. If the SIM was swapped recently, or the device is outside the usual store area, the agent blocks upload and export of financial files until network identity is confirmed, without relying on SMS alone."
)

add_sub("Expected benefits")
add_bullet(
    "For the merchant: real net profit (not phantom product profit), daily decisions, and protection of profit files against a stolen SIM or an unusual location."
)
add_bullet(
    "For operators and GSMA Open Gateway: a real B2B FinTech that consumes CAMARA APIs, not a slide-only integration."
)
add_bullet(
    "For MENA startups: a SaaS model that can scale across Arabic-first SMEs that still live in Excel."
)
add_bullet(
    "For digital inclusion: Arabic RTL and English so independent merchants, not only English-speaking corporates, can use secure FinTech."
)

add_heading_bar("4. Project type and alignment with GSMA pillars")
add_body(
    "Project type: B2B SaaS FinTech for small and medium merchants, with a programmable-network security layer and an AI Agent that orchestrates CAMARA APIs."
)
add_body("Alignment with GSMA Open Gateway and digital development:")
add_bullet(
    "Digital transformation of SMEs: replaces scattered Excel with structured P&L, diagnosis and bilingual advice — the digitalisation GSMA seeks in the regional economy."
)
add_bullet(
    "Startup enablement: a shippable product architecture (Next.js / TypeScript / PostgreSQL) that a MENA startup can commercialise, not a one-off demo script."
)
add_bullet(
    "Open Gateway / CAMARA: identity and anti-fraud are delivered from the mobile network (4G/5G), which is the GSMA Open Gateway thesis — network capabilities as APIs."
)
add_bullet(
    "Nokia Network-as-Code: one developer access path to operator CAMARA implementations, which is how GSMA and Nokia ask teams to build."
)
add_bullet(
    "Regional inclusion: built from Gaza for MENA merchants, with Arabic RTL and English, so secure FinTech is not limited to large corporates."
)

add_heading_bar("5. Theme")
kv_table(
    [
        (
            "Theme selected",
            "Theme 4: Secure FinTech, Payments & Anti-Fraud Innovation",
        ),
        (
            "Why this theme",
            "The product protects merchant financial accounts and files against SIM-swap takeover, OTP abuse, and uploads from untrusted locations, while remaining a payments-adjacent FinTech (P&L, pricing, cash decisions).",
        ),
    ]
)
add_body(
    "This is Theme 4 by design: anti-fraud and identity from the network, applied to the merchant's money files, orchestrated by an AI Agent."
)

add_heading_bar("6. API Usage (Nokia Network-as-Code / CAMARA)")
add_body(
    "All network calls go through the Nokia Network-as-Code platform, which hosts GSMA Open Gateway CAMARA APIs. The Smart Guard AI Agent orchestrates the calls and takes the decision."
)
add_body(
    "Testing and prototype validation will utilize Nokia Network-as-Code developer portal simulators for SIM Swap, Location Verification, and Number Verification. This confirms that the freeze policy can be demonstrated during the hackathon without waiting for a live operator contract, then mapped one-to-one onto production CAMARA endpoints."
)

api_rows = [
    ("Nokia / CAMARA API", "When called", "Agent decision"),
    (
        "Number Verification",
        "Login and step-up identity: a silent 4G/5G check, with no SMS.",
        "Confirm that the merchant is on the claimed mobile number from the network itself, not from a spoofed message. Do not trust SMS alone.",
    ),
    (
        "SIM Swap",
        "Login, password reset, price changes, and other sensitive financial actions.",
        "If the SIM was swapped recently, freeze the session. SMS OTP is treated as untrusted.",
    ),
    (
        "Location Verification",
        "Upload of Excel/CSV/PDF financial files and export of P&L reports.",
        "The device must match the merchant's usual store or branch area. Otherwise the upload is blocked.",
    ),
]
t = doc.add_table(rows=len(api_rows), cols=3)
for i, row in enumerate(api_rows):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        if i == 0:
            set_run(run, 9, True, WHITE)
            shade_cell(cell, "0B1120")
        else:
            set_run(run, 9, False, BLACK)
            shade_cell(cell, "F3F6FB" if i % 2 == 0 else "FFFFFF")
        set_cell_border(cell)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
set_run(p.add_run("Mandatory AI Agent (requirement 2): "), 11, True, NAVY)
set_run(
    p.add_run(
        "Smart Guard is not a caption generator. It calls Number Verification, SIM Swap and Location Verification as tools, combines them with financial context, and executes Allow / Step-up / Freeze. Example: if SIM Swap is recent OR location is outside the store during an Excel upload, the workbook is not ingested until Number Verification succeeds on the trusted line."
    ),
    11,
    False,
    BLACK,
)

add_sub("Declaration")
add_body(
    "This idea uses GSMA CAMARA network APIs through Nokia Network-as-Code, and uses an AI Agent to orchestrate those APIs and take security decisions. Both compulsory hackathon conditions are met. Prototype testing will run on Nokia Network-as-Code developer portal simulators."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(
    p.add_run(
        "Submitted by Israa Nael Hamad  ·  Software Engineering, University of Palestine (Gaza)  ·  Smart Profits Team  ·  14 August 2026"
    ),
    10,
    True,
    MUTED,
)

out_dir = Path(__file__).resolve().parent
primary = out_dir / "Smart-Profits-IDEA-CAPTURE-TEMPLATE.docx"
fallback = out_dir / "Smart-Profits-IDEA-CAPTURE-TEMPLATE-EN.docx"
try:
    doc.save(primary)
    print("saved", primary)
except PermissionError:
    doc.save(fallback)
    print("locked original; saved", fallback)
